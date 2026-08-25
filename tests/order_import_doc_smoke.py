from __future__ import annotations

import sys
import shutil
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import order_system.order_import as order_import  # noqa: E402


PNG_BYTES = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
    b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00"
    b"\x00\x0cIDATx\x9cc``\x00\x00\x00\x04\x00\x01\xf6"
    b"\x178U\x00\x00\x00\x00IEND\xaeB`\x82"
)


root = Path(tempfile.mkdtemp(prefix="twd-import-doc-"))
legacy_doc = root / "customer-order.doc"
legacy_doc.write_bytes(b"legacy-word-placeholder")

original_convert = order_import._convert_with_libreoffice
original_contains_images = order_import._docx_contains_images
original_extract_docx = order_import._extract_docx
original_render = order_import._render_doc_pages
original_extract_document_text = order_import.extract_document_text
original_layout_render = order_import._render_layout_document_pages


def placeholder_docx(output_dir: Path) -> Path:
    target = output_dir / "customer-order.docx"
    target.write_bytes(b"docx-placeholder")
    return target


try:
    def fake_text_convert(source_path: Path, output_format: str, output_dir: Path) -> Path:
        assert source_path == legacy_doc
        assert output_format == "docx"
        return placeholder_docx(output_dir)

    order_import._convert_with_libreoffice = fake_text_convert
    order_import._extract_docx = lambda path: "产品：钥匙扣\n[表格1]\n数量 | 100"
    text_content = order_import._legacy_doc_user_content(legacy_doc, "数量以表格为准")
    assert isinstance(text_content, str)
    assert "产品：钥匙扣" in text_content
    assert "数量 | 100" in text_content
    assert "数量以表格为准" in text_content

    def fake_image_convert(source_path: Path, output_format: str, output_dir: Path) -> Path:
        assert output_format == "docx"
        return placeholder_docx(output_dir)

    def fake_render(source_path: Path, output_dir: Path) -> list[Path]:
        image_path = output_dir / "doc-page-1.png"
        image_path.write_bytes(PNG_BYTES)
        return [image_path]

    order_import._convert_with_libreoffice = fake_image_convert
    order_import._extract_docx = lambda path: ""
    order_import._docx_contains_images = lambda path: True
    order_import._render_doc_pages = fake_render
    visual_content = order_import._legacy_doc_user_content(legacy_doc, "")
    assert isinstance(visual_content, list)
    assert visual_content[0]["type"] == "text"
    assert visual_content[1]["type"] == "image_url"
    assert visual_content[1]["image_url"]["url"].startswith("data:image/png;base64,")

    layout_order = root / "customer-order.xlsx"
    layout_order.write_bytes(b"xlsx-placeholder")

    def fake_layout_render(source_path: Path, output_dir: Path) -> list[Path]:
        assert source_path == layout_order
        image_path = output_dir / "doc-page-1.png"
        image_path.write_bytes(PNG_BYTES)
        return [image_path]

    order_import.extract_document_text = lambda path: "材质 | 锌合金烤漆\n电镀 | 染黑"
    order_import._render_layout_document_pages = fake_layout_render
    layout_content = order_import._layout_document_user_content(layout_order, "红色方块为选中")
    assert isinstance(layout_content, list)
    assert "锌合金烤漆" in layout_content[0]["text"]
    assert "红色方块为选中" in layout_content[0]["text"]
    assert layout_content[1]["min_pixels"] == order_import.VISUAL_MIN_PIXELS

    def failed_layout_render(source_path: Path, output_dir: Path) -> list[Path]:
        raise order_import.OrderImportError("render failed")

    order_import._render_layout_document_pages = failed_layout_render
    fallback_content = order_import._layout_document_user_content(layout_order, "")
    assert isinstance(fallback_content, str)
    assert "锌合金烤漆" in fallback_content
finally:
    order_import._convert_with_libreoffice = original_convert
    order_import._docx_contains_images = original_contains_images
    order_import._extract_docx = original_extract_docx
    order_import._render_doc_pages = original_render
    order_import.extract_document_text = original_extract_document_text
    order_import._render_layout_document_pages = original_layout_render

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Border, Side

spreadsheet_path = root / "structured-order.xlsx"
spreadsheet = Workbook()
sheet = spreadsheet.active
sheet.title = "Sheet1"
sheet.merge_cells("A1:H1")
sheet["A1"] = "测试订单"
sheet["A5"] = "产品品名"
sheet["C5"] = "描述"
sheet["A6"] = "证章"
sheet["C6"] = "低温锌合金，最大尺寸20mm，电镀镍，三面抛，不入色，配件蝴夹x1pc"
sheet["A11"] = "备注"
sheet["H11"] = "测试"
sheet["M6"].border = Border(left=Side(style="thin"))
spreadsheet.save(spreadsheet_path)

prepared_path = order_import._prepare_spreadsheet_for_render(spreadsheet_path, root / "prepared")
assert prepared_path != spreadsheet_path and prepared_path.is_file()
prepared = load_workbook(prepared_path)
prepared_sheet = prepared["Sheet1"]
assert order_import._worksheet_visual_bounds(prepared_sheet) == (1, 1, 11, 8)
assert str(prepared_sheet.print_area) == "'Sheet1'!$A$1:$H$11"
assert prepared_sheet.page_setup.fitToWidth == 1
assert prepared_sheet.page_setup.fitToHeight == 0
assert prepared_sheet.page_setup.scale is None
assert prepared_sheet.sheet_properties.pageSetUpPr.fitToPage is True
prepared.close()

assert ".doc" in order_import.SUPPORTED_DOCUMENT_SUFFIXES

if shutil.which("libreoffice") and shutil.which("pdftoppm"):
    from docx import Document

    real_root = root / "real-conversion"
    real_root.mkdir()
    source_docx = real_root / "real-order.docx"
    source = Document()
    source.add_paragraph("产品：钥匙扣")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "数量"
    table.cell(0, 1).text = "100"
    source.save(source_docx)

    legacy_path = order_import._convert_with_libreoffice(source_docx, "doc", real_root / "to-doc")
    extracted = order_import.extract_document_text(legacy_path)
    assert "产品：钥匙扣" in extracted
    assert "数量" in extracted
    assert "100" in extracted

    page_paths = order_import._render_doc_pages(legacy_path, real_root / "rendered")
    assert page_paths and all(path.stat().st_size > 0 for path in page_paths)

print(f"order import doc smoke ok: {root}")
